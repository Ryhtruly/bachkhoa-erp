import os
from docx import Document

TEMPLATE_PATH = os.path.join(os.getcwd(), "src", "templates", "mau_hop_dong.docx")
OUTPUT_DIR = os.path.join(os.getcwd(), "src", "static", "generated_contracts")

def generate_contract(data):
    """
    Generate a contract .docx file by replacing placeholders in the template.
    data: dictionary containing placeholder values (e.g., SO_HOP_DONG, TEN_KHACH_HANG)
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    try:
        # Load the template
        doc = Document(TEMPLATE_PATH)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, data)
            
        # Replace placeholders in tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, data)
                        
        # Save output file
        filename = f"HopDong_{data.get('MA_HO_SO', 'TEMP')}_{data.get('SO_HOP_DONG', 'TEMP')}.docx"
        filename = filename.replace("/", "_") # Sanitize filename
        output_path = os.path.join(OUTPUT_DIR, filename)
        doc.save(output_path)
        
        # Return web-accessible path
        return True, f"/static/generated_contracts/{filename}", output_path
    except Exception as e:
        print(f"Error generating contract: {e}")
        return False, str(e), None

def replace_placeholders_in_paragraph(paragraph, data):
    """Helper function to replace placeholder text while preserving formatting."""
    # Build complete paragraph text
    full_text = "".join(run.text for run in paragraph.runs)
    
    # Check if any placeholder exists in the paragraph text
    has_placeholder = False
    for key in data.keys():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            has_placeholder = True
            break
            
    if not has_placeholder:
        return
        
    # Replace in full text
    for key, val in data.items():
        placeholder = f"{{{{{key}}}}}"
        if val is None:
            val = ""
        full_text = full_text.replace(placeholder, str(val))
        
    # Clear runs and write updated text in a single run to prevent formatting breakages
    if len(paragraph.runs) > 0:
        # Save style of first run if available
        first_run = paragraph.runs[0]
        bold = first_run.bold
        italic = first_run.italic
        font_name = first_run.font.name if first_run.font else None
        font_size = first_run.font.size if first_run.font else None
        
        # Clear all runs
        paragraph.text = ""
        
        # Add new single run
        new_run = paragraph.add_run(full_text)
        new_run.bold = bold
        new_run.italic = italic
        if font_name:
            new_run.font.name = font_name
        if font_size:
            new_run.font.size = font_size
    else:
        paragraph.text = full_text
